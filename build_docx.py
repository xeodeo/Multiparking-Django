#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Genera Documentacion_Multiparking_COMPLETA.docx con estilos idénticos a V4
(Heading 1-4 + Normal, sin TITLE/SUBTITLE) y tablas reales para RF, RNF,
Diccionario de Datos, Encuesta, Riesgos y Pruebas.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── helpers ──────────────────────────────────────────────────────────────────
def h1(text):
    doc.add_heading(text, level=1)

def h2(text):
    doc.add_heading(text, level=2)

def h3(text):
    doc.add_heading(text, level=3)

def h4(text):
    doc.add_heading(text, level=4)

def p(text='', bold=False):
    para = doc.add_paragraph()
    if text:
        run = para.add_run(text)
        run.bold = bold
    return para

def pb(text):
    """Párrafo con bold"""
    return p(text, bold=True)

def tabla_rf(codigo, nombre, descripcion, entradas, prioridad):
    """Tabla de 5 filas × 2 cols — formato V4"""
    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = 'Table Grid'
    rows_data = [
        ('Código Requisito', codigo),
        ('Nombre', nombre),
        ('Descripción', descripcion),
        ('Entradas', entradas),
        ('Prioridad', prioridad),
    ]
    for i, (k, v) in enumerate(rows_data):
        row = tbl.rows[i]
        row.cells[0].text = k
        row.cells[1].text = v
        # Bold the key column
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True
    doc.add_paragraph()  # spacing

def tabla_generica(headers, rows, titulo=None):
    """Tabla genérica con encabezados en negrita"""
    if titulo:
        p(titulo, bold=True)
    cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=cols)
    tbl.style = 'Table Grid'
    # Header row
    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        hrow.cells[i].text = h
        for run in hrow.cells[i].paragraphs[0].runs:
            run.bold = True
    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            row.cells[c_idx].text = str(val)
    doc.add_paragraph()  # spacing

# ═══════════════════════════════════════════════════════════════════════════
# PORTADA  (todo Normal — igual que V4)
# ═══════════════════════════════════════════════════════════════════════════
p('Sistema de Información Web para la Gestión Integral de Parqueaderos Privados', bold=True)
p('MultiParking', bold=True)
p()
p('Integrantes:')
p('Kevin Arley Grisales Ovalle\nKevin Santiago Pinto López')
p()
p('Análisis y Desarrollo de Software')
p('Ficha: 3119175')
p()
p('Centro de Diseño y Metrología – SENA')
p('Bogotá D.C., 2026')
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# RESUMEN / ABSTRACT
# ═══════════════════════════════════════════════════════════════════════════
h1('Resumen')
p(
    'MultiParking es un sistema de información web orientado a la automatización y gestión integral '
    'de parqueaderos privados. La plataforma cubre el registro de usuarios con tres roles diferenciados '
    '(administrador, vigilante y cliente), el control de vehículos propios y visitantes, la asignación '
    'y liberación de espacios en tiempo real, el cálculo automático de tarifas por minuto con redondeo '
    'a cien pesos colombianos, la generación de pagos, la aplicación de cupones de descuento, la gestión '
    'de reservas, el registro de novedades operativas con notificación por correo electrónico, un sistema '
    'de fidelización mediante stickers acumulables y la generación de reportes exportables. El sistema se '
    'desarrolló en Python con Django 6.0, PostgreSQL como gestor de base de datos en producción (Render.com), '
    'Tailwind CSS para el diseño responsivo, SendGrid para correos transaccionales y Gunicorn como servidor WSGI.'
)
p('Palabras clave: parqueadero, sistema de información web, Django, PostgreSQL, gestión de espacios, fidelización.')
p()

h1('Abstract')
p(
    'MultiParking is a web-based information system designed for the comprehensive automation and management '
    'of private parking lots. The platform covers user registration with three differentiated roles '
    '(administrator, security guard, and client), real-time space assignment, automatic rate calculation, '
    'payment generation, coupon application, reservation management, operational incident logging with email '
    'notification, a loyalty sticker system, and exportable reports. Built with Python/Django 6.0, '
    'PostgreSQL (Render.com production), Tailwind CSS, SendGrid and Gunicorn.'
)
p('Keywords: parking lot, web information system, Django, PostgreSQL, space management, loyalty.')
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 1. INTRODUCCIÓN
# ═══════════════════════════════════════════════════════════════════════════
h1('Introducción')
p(
    'En la actualidad, los parqueaderos privados de pequeña y mediana escala operan con métodos manuales '
    'que generan ineficiencias operativas: registros en papel, cálculos de tarifa imprecisos, ausencia de '
    'trazabilidad y nula visibilidad del estado de los espacios en tiempo real. Esta situación deriva en '
    'pérdida de ingresos, inconformidad de los usuarios y dificultad para tomar decisiones basadas en datos '
    '(García & Martínez, 2022).'
)
p(
    'El proyecto MultiParking desarrolla una solución web que automatiza de forma integral la operación de '
    'parqueaderos privados. El sistema registra ingresos y salidas de vehículos, calcula tarifas en tiempo '
    'real, gestiona reservas anticipadas, emite notificaciones digitales, aplica descuentos mediante cupones '
    'y acumula beneficios de fidelización.'
)
p(
    'La arquitectura tecnológica se apoya en Django 6.0 (Python 3.12), PostgreSQL en producción (Render.com), '
    'Tailwind CSS vía CDN, Chart.js para visualizaciones estadísticas, SendGrid como proveedor de correo '
    'transaccional y Gunicorn como servidor WSGI. La autenticación implementa un modelo de sesión propio '
    '—sin dependencia del módulo django.contrib.auth.User— con tres roles: ADMIN, VIGILANTE y CLIENTE.'
)
p(
    'Este documento sigue la estructura definida por el programa Análisis y Desarrollo de Software (ADSO) '
    'del SENA y las normas de presentación APA séptima edición (American Psychological Association, 2020).'
)

# ═══════════════════════════════════════════════════════════════════════════
# 2. PLANTEAMIENTO DEL PROBLEMA
# ═══════════════════════════════════════════════════════════════════════════
h1('Planteamiento del Problema')
h2('Árbol del Problema')
h3('Efectos')
pb('Pérdida de información y trazabilidad en el registro de vehículos:')
p('causada por registros manuales dispersos y falta de integración del sistema.')
pb('Cobros imprecisos o injustos por uso del parqueadero:')
p('debido a la ausencia de tarifación confiable y auditoría en tiempo real.')
pb('Filas y demoras en el ingreso y salida de vehículos:')
p('por control manual y verificación lenta.')
pb('Inconformidad de los residentes y visitantes:')
p('generada por tiempos de espera, cobros erróneos y baja transparencia operativa.')

h3('Problema Principal')
pb(
    'Ineficiencia, falta de trazabilidad y pobre control en parqueaderos privados, '
    'ocasionado por procesos manuales, información no centralizada y carencia de '
    'herramientas digitales que soporten la operación, el cobro y la toma de decisiones '
    '(Bookings & Hernández, 2023).'
)

h3('Causas')
pb('Control manual del ingreso y salida de vehículos:')
p('provoca cuellos de botella y errores humanos.')
pb('Registro de información en papel o planillas físicas:')
p('genera pérdida de datos y dificultad para auditorías.')
pb('Falta de infraestructura tecnológica:')
p('limita la automatización y trazabilidad.')
pb('Ausencia de herramientas digitales para cobros y reservas:')
p('impide calcular tarifas correctas y agilizar pagos.')

h2('Descripción del Problema')
p(
    'La mayoría de los parqueaderos privados realizan sus procesos operativos y administrativos de manera '
    'manual. El control del ingreso y salida de vehículos depende de planillas físicas, lo que impide tener '
    'trazabilidad confiable y genera inconsistencias en los tiempos y valores cobrados. La ausencia de un '
    'sistema tecnológico centralizado limita la visibilidad en tiempo real de los espacios disponibles, '
    'dificulta la conciliación de pagos y complica la detección de errores o fraudes. La imposibilidad de '
    'reservar espacio anticipadamente reduce la satisfacción del usuario y la competitividad del establecimiento '
    'frente a alternativas digitales.'
)

# ═══════════════════════════════════════════════════════════════════════════
# 3. OBJETIVOS
# ═══════════════════════════════════════════════════════════════════════════
h1('Objetivos')
h2('Objetivo General')
p(
    'Desarrollar un sistema de información web que permita la gestión integral de parqueaderos privados, '
    'optimizando el control de ingreso y salida de vehículos, la asignación de espacios, el cálculo '
    'automático de tarifas, la gestión de reservas anticipadas y el seguimiento de novedades operativas.'
)

h2('Objetivos Específicos')
pb('Fase de inicio:')
p('Analizar los procesos actuales de operación y control en parqueaderos privados para identificar los requerimientos funcionales y no funcionales del sistema.')
p('Identificar las necesidades de los usuarios (administradores, vigilantes y clientes) mediante encuestas y observación directa.')

pb('Fase de planificación:')
p('Diseñar la arquitectura del sistema web, definiendo la estructura de la base de datos, los módulos funcionales y la interacción entre componentes.')
p('Elaborar los diagramas de casos de uso, clases y flujo de datos que representen el comportamiento del sistema.')

h4('Fase de ejecución:')
p('Desarrollar los once módulos principales: usuarios, vehículos, pisos/espacios, inventario de parqueos, tarifas, pagos, cupones, reservas, novedades, fidelización y reportes.')
p('Implementar el cálculo automático de tarifas, las notificaciones por correo electrónico y la exportación de reportes en PDF y Excel.')

h4('Fase de implementación:')
p('Desplegar el sistema en Render.com, validando su correcto funcionamiento con datos reales.')
p('Capacitar al personal y realizar pruebas piloto antes de la puesta en marcha final.')

# ═══════════════════════════════════════════════════════════════════════════
# 4. JUSTIFICACIÓN
# ═══════════════════════════════════════════════════════════════════════════
h1('Justificación')
p(
    'En numerosos parqueaderos privados de pequeña y mediana escala, el control manual de ingreso y salida '
    'de vehículos genera demoras, fallos en el cobro y pérdida de información crítica. Estudios recientes '
    'señalan que la adopción de sistemas digitales de gestión de aparcamientos mejora la eficiencia operativa, '
    'reduce errores humanos y ofrece trazabilidad de datos en tiempo real (García & Martínez, 2022). '
    'Así mismo, Garavito Albao et al. (2021) demuestran que los sistemas de información en pequeñas empresas '
    'de servicios reducen el tiempo de atención hasta en un 40% y disminuyen los errores de registro en '
    'más del 60%.'
)
p(
    'La plataforma beneficia directamente a tres perfiles de usuarios: administradores, que obtienen '
    'visibilidad total y reportes de inteligencia operativa; vigilantes, que agilizan los procesos de '
    'ingreso y salida; y clientes, que acceden a reservas anticipadas, historial de parqueos y beneficios '
    'de fidelización. La implementación sobre Django 6.0 garantiza mantenibilidad a largo plazo gracias a '
    'su arquitectura MVT y su comunidad activa.'
)

# ═══════════════════════════════════════════════════════════════════════════
# 5. ALCANCE
# ═══════════════════════════════════════════════════════════════════════════
h1('Alcance del Proyecto')
p(
    'El proyecto consiste en desarrollar un sistema de información web denominado MultiParking, orientado '
    'a automatizar la gestión de parqueaderos privados de pequeña y mediana escala. El sistema opera como '
    'plataforma web responsiva desde dispositivos móviles o computadores sin requerir hardware especializado.'
)
p(
    'El alcance incluye: levantamiento de requerimientos; diseño de la arquitectura y la base de datos '
    'relacional (PostgreSQL); once módulos funcionales; notificaciones por correo (SendGrid); exportación '
    'a PDF (reportlab) y Excel (openpyxl); despliegue en Render.com.'
)
p(
    'Quedan fuera del alcance: integración con hardware automatizado (barreras, sensores, cámaras) y '
    'desarrollo de aplicación móvil nativa, contemplados para versiones futuras.'
)

# ═══════════════════════════════════════════════════════════════════════════
# 6. MATRIZ DE RIESGO
# ═══════════════════════════════════════════════════════════════════════════
h1('Matriz de Riesgo')
p('La Tabla 1 presenta los riesgos identificados durante la planificación del proyecto.')
p()
tabla_generica(
    ['Nº', 'Riesgo', 'Causa', 'Impacto', 'Probabilidad', 'Estrategia de mitigación'],
    [
        ['1', 'Retrasos en el desarrollo del sistema', 'Subestimación de tiempos, problemas técnicos', 'Aplazamiento en la entrega final', 'Media', 'Cronograma ágil con revisión semanal.'],
        ['2', 'Dificultades en la integración de pagos digitales', 'Problemas con APIs de pasarelas de pago', 'No se podrán realizar cobros automatizados', 'Media', 'Usar APIs documentadas como PayU o Wompi.'],
        ['3', 'Falta de conectividad en el entorno piloto', 'Infraestructura deficiente de red', 'Fallos en el acceso o carga del sistema', 'Alta', 'Verificar requisitos técnicos del sitio previamente.'],
        ['4', 'Baja adopción por parte de los usuarios', 'Resistencia al cambio', 'Uso limitado del sistema', 'Media', 'Capacitaciones y pruebas piloto con retroalimentación.'],
        ['5', 'Pérdida de datos durante el uso', 'Fallos en el servidor', 'Pérdida de historial y registros', 'Baja', 'Implementar backups automáticos diarios.'],
        ['6', 'Cambios en los requerimientos del cliente', 'Nuevas necesidades descubiertas', 'Re-trabajo o rediseño de módulos', 'Media', 'Reuniones frecuentes con stakeholders.'],
        ['7', 'Salida de un integrante clave del equipo', 'Problemas personales o deserción', 'Retrasos o redistribución de tareas', 'Baja', 'Documentación compartida y roles cruzados.'],
    ],
    titulo='Tabla 1\nMatriz de Riesgos del Proyecto MultiParking'
)
p('Nota. Elaboración propia (2026).')

# ═══════════════════════════════════════════════════════════════════════════
# 7. ELICITACIÓN DE REQUISITOS
# ═══════════════════════════════════════════════════════════════════════════
h1('Elicitación de Requisitos')
h2('Identificación de Procesos')
p('Los procesos principales identificados en parqueaderos privados son:')
p('• Registro de ingreso de vehículo con asignación automática de espacio disponible.')
p('• Cálculo de tiempo de permanencia y tarifa al momento de la salida.')
p('• Gestión de pagos y aplicación de descuentos mediante cupones.')
p('• Reserva anticipada de espacios por parte de clientes registrados.')
p('• Registro de novedades operativas con notificación al usuario afectado.')
p('• Generación de reportes de ingresos, ocupación y actividad del parqueadero.')

h2('Recolección de la Información del Software a Construir')
p(
    'Para iniciar el desarrollo del sistema MultiParking se realizó una fase de recolección de datos '
    'con el objetivo de comprender las necesidades reales de propietarios, vigilantes y usuarios. '
    'La información se obtuvo a través de dos técnicas principales: observación directa en parqueaderos '
    'locales de la localidad de Suba, barrio Compartir, y encuestas digitales dirigidas a usuarios.'
)

h2('Elección de la Técnica de Recolección de la Información')
p('Para el proyecto se optó por utilizar las siguientes técnicas:')
p('Encuesta digital: Aplicada a 13 usuarios y administradores de parqueaderos, con 13 preguntas de selección múltiple sobre su experiencia actual y expectativas frente a un sistema tecnológico.')
p('Observación directa: Empleada para registrar de forma sistemática el comportamiento de los procesos operativos en parqueaderos de pequeño y mediano tamaño.')

h2('Aplicación de la Técnica de Recolección de Información')
p('La encuesta se aplicó de forma digital. La observación se efectuó en parqueaderos reales de la localidad de Suba en días de alta y baja afluencia.')

h2('Organización de la Información Recolectada')
h3('Encuesta')
p('A continuación se presenta el análisis estadístico de la información recolectada en la encuesta (n = 13):')

# Survey tables
surveys = [
    ('¿Con qué frecuencia utilizas parqueaderos públicos o privados?',
     [('Todos los días', '4', '4', '0.308', '0.308'),
      ('Varias veces a la semana', '3', '7', '0.231', '0.539'),
      ('Ocasionalmente', '3', '10', '0.231', '0.770'),
      ('Casi nunca', '3', '13', '0.231', '1.000'),
      ('Muestra', '13', '', '1', '')]),
    ('¿Qué medio usas actualmente para pagar en un parqueadero?',
     [('Efectivo', '11', '11', '0.846', '0.846'),
      ('Tarjeta débito/crédito', '1', '12', '0.077', '0.923'),
      ('App o sistema digital', '1', '13', '0.077', '1.000'),
      ('Muestra', '13', '', '1', '')]),
    ('¿Qué aspectos te resultan más molestos al usar un parqueadero?',
     [('Pago solo en efectivo', '7', '7', '0.538', '0.538'),
      ('Largos tiempos de espera', '6', '13', '0.462', '1.000'),
      ('Tarifas poco claras', '6', '19', '0.462', '1.462'),
      ('Filas para entrar o salir', '6', '25', '0.462', '1.923'),
      ('Falta de espacios disponibles', '5', '30', '0.385', '2.308'),
      ('Pérdida de tiquetes', '2', '32', '0.154', '2.462'),
      ('Muestra', '13', '', '', '')]),
    ('¿Te gustaría poder ingresar al parqueadero escaneando un código QR?',
     [('Sí', '5', '5', '0.385', '0.385'),
      ('Tal vez', '8', '13', '0.615', '1.000'),
      ('No', '0', '13', '0', '1.000'),
      ('Muestra', '13', '', '1', '')]),
    ('¿Te parecería útil poder ver el tiempo y costo acumulado en tiempo real?',
     [('Sí, muy útil', '12', '12', '0.923', '0.923'),
      ('Me es indiferente', '0', '12', '0', '0.923'),
      ('No lo necesito', '1', '13', '0.077', '1.000'),
      ('Muestra', '13', '', '1', '')]),
    ('¿Usarías una app para reservar espacio de parqueo antes de llegar?',
     [('Sí', '8', '8', '0.615', '0.615'),
      ('Tal vez', '4', '12', '0.308', '0.923'),
      ('No', '1', '13', '0.077', '1.000'),
      ('Muestra', '13', '', '1', '')]),
    ('¿Qué métodos de pago preferirías para este sistema?',
     [('Transferencia digital / billetera', '10', '10', '0.769', '0.769'),
      ('Tarjeta débito/crédito', '8', '18', '0.615', '1.385'),
      ('Efectivo', '6', '24', '0.462', '1.846'),
      ('Pago por código QR', '5', '29', '0.385', '2.231'),
      ('Muestra', '13', '', '', '')]),
    ('¿Qué tanto te gustaría que el parqueadero reconociera tu historial para darte descuentos?',
     [('Me gustaría mucho', '6', '6', '0.462', '0.462'),
      ('Es algo interesante', '5', '11', '0.385', '0.847'),
      ('No me interesa', '2', '13', '0.154', '1.000'),
      ('Muestra', '13', '', '1', '')]),
    ('¿Consideras importante salir del parqueadero rápidamente sin interactuar con personal?',
     [('Moderadamente importante', '6', '6', '0.462', '0.462'),
      ('Muy importante', '5', '11', '0.385', '0.847'),
      ('No es importante', '2', '13', '0.154', '1.000'),
      ('Muestra', '13', '', '1', '')]),
    ('¿Qué tan dispuesto estarías a usar un sistema de parqueadero completamente automatizado?',
     [('Muy dispuesto', '7', '7', '0.538', '0.538'),
      ('Algo dispuesto', '6', '13', '0.462', '1.000'),
      ('No dispuesto', '0', '13', '0', '1.000'),
      ('Muestra', '13', '', '1', '')]),
    ('¿Estarías dispuesto/a a pagar un poco más por un parqueadero automatizado?',
     [('Solo si el precio es igual al tradicional', '7', '7', '0.538', '0.538'),
      ('Sí', '6', '13', '0.462', '1.000'),
      ('No', '0', '13', '0', '1.000'),
      ('Muestra', '13', '', '1', '')]),
    ('¿Qué funcionalidad te gustaría que tuviera un parqueadero automatizado?',
     [('Poder reservar', '3', '3', '0.231', '0.231'),
      ('Pagar al ingresar, no al salir', '2', '5', '0.154', '0.385'),
      ('Garantizar seguridad del vehículo', '3', '8', '0.231', '0.612'),
      ('Notificación del vencimiento del parqueo', '3', '11', '0.231', '0.846'),
      ('Parqueadero más amplio', '2', '13', '0.154', '1.000'),
      ('Muestra', '13', '', '1', '')]),
]

headers_survey = ['Variable', 'Frecuencia Absoluta', 'Frecuencia Absoluta Acumulada', 'Frecuencia Relativa', 'Frecuencia Relativa Acumulada']
for i, (pregunta, filas) in enumerate(surveys, 1):
    h3(pregunta)
    tabla_generica(headers_survey, filas)
    p(f'Nota. Elaboración propia (2026).')
    p()

h3('Observación')
p(
    'Se llevaron a cabo observaciones directas en parqueaderos privados de tamaño pequeño y mediano, '
    'en días con gran y poca afluencia. Se registraron los procedimientos de entrada y salida, la '
    'asignación de espacios, los tiempos de espera, las formas de cobro y las incidencias más frecuentes. '
    'Los hallazgos confirmaron la necesidad de digitalizar el control de acceso, el cobro y la '
    'comunicación con los usuarios.'
)

# ═══════════════════════════════════════════════════════════════════════════
# 8. REQUERIMIENTOS
# ═══════════════════════════════════════════════════════════════════════════
h1('Requerimientos')

h2('Módulo Usuarios')
rfs_usuarios = [
    ('RF01', 'Registrar usuario', 'El sistema permitirá registrar un nuevo usuario capturando nombre, apellido, documento de identidad, correo electrónico, teléfono y contraseña. La contraseña se almacenará cifrada con PBKDF2-SHA256.', 'usuNombre, usuApellido, usuDocumento, usuCorreo, usuTelefono, usuClaveHash', 'Alta'),
    ('RF02', 'Iniciar sesión', 'El sistema permitirá autenticarse mediante correo y contraseña. Se verificará el hash y se creará sesión con datos del usuario y su rol.', 'usuCorreo, contraseña', 'Alta'),
    ('RF03', 'Cerrar sesión', 'El sistema permitirá finalizar la sesión activa, eliminando los datos de sesión del servidor.', 'Sesión activa', 'Alta'),
    ('RF04', 'Recuperar contraseña', 'El sistema enviará al correo registrado un enlace de recuperación con token de un solo uso. Al usarlo, el token queda invalidado automáticamente.', 'usuCorreo', 'Alta'),
    ('RF05', 'Gestionar usuarios (Admin)', 'El administrador podrá crear, editar, activar/desactivar y eliminar usuarios del sistema.', 'usuNombre, usuApellido, usuDocumento, usuCorreo, usuTelefono, rolTipoRol', 'Alta'),
    ('RF06', 'Asignar rol', 'El sistema soportará tres roles: ADMIN, VIGILANTE y CLIENTE, con acceso diferenciado.', 'rolTipoRol (ENUM)', 'Alta'),
    ('RF07', 'Ver perfil propio', 'El usuario podrá consultar sus datos personales, historial de parqueos, stickers acumulados y reservas activas.', 'Sesión activa', 'Media'),
    ('RF08', 'Actualizar datos personales', 'El usuario podrá modificar su nombre, apellido, teléfono y contraseña desde su perfil.', 'usuNombre, usuApellido, usuTelefono, nueva contraseña', 'Media'),
]
for args in rfs_usuarios:
    tabla_rf(*args)

h2('Módulo Vehículos')
rfs_vehiculos = [
    ('RF09', 'Registrar vehículo propio', 'El cliente podrá registrar vehículos asociados a su cuenta indicando placa, tipo, color, marca y modelo.', 'vehPlaca, vehTipo, vehColor, vehMarca, vehModelo, fkIdUsuario', 'Alta'),
    ('RF10', 'Registrar vehículo visitante', 'El vigilante podrá registrar vehículos de visitantes sin cuenta, capturando placa, tipo, nombre y teléfono de contacto.', 'vehPlaca, vehTipo, nombre_contacto, telefono_contacto', 'Alta'),
    ('RF11', 'Listar vehículos por usuario', 'El sistema mostrará la lista de vehículos registrados por el usuario autenticado.', 'fkIdUsuario (sesión)', 'Media'),
    ('RF12', 'Editar vehículo', 'El usuario podrá modificar los datos de un vehículo propio.', 'vehPlaca, vehTipo, vehColor, vehMarca, vehModelo', 'Media'),
    ('RF13', 'Eliminar vehículo', 'El administrador podrá eliminar un vehículo si no tiene registros activos ni reservas vigentes.', 'idVehiculo', 'Baja'),
]
for args in rfs_vehiculos:
    tabla_rf(*args)

h2('Módulo Espacios')
rfs_espacios = [
    ('RF14', 'Gestionar pisos', 'El administrador podrá crear, editar y desactivar pisos del parqueadero. No se elimina un piso con espacios ocupados.', 'pisNombre, pisEstado', 'Alta'),
    ('RF15', 'Gestionar tipos de espacio', 'El administrador podrá crear y editar tipos de espacio (CARRO, MOTO, BICICLETA).', 'nombre (TipoEspacio)', 'Alta'),
    ('RF16', 'Gestionar espacios', 'El administrador podrá crear, editar y eliminar espacios individuales por piso y tipo.', 'espNumero, fkIdPiso, fkIdTipoEspacio, espEstado', 'Alta'),
    ('RF17', 'Ver disponibilidad en tiempo real', 'El sistema mostrará la vista general de pisos coloreada por estado (DISPONIBLE/OCUPADO/RESERVADO/MANTENIMIENTO), actualizada cada 30 segundos vía AJAX.', 'Ninguna (dashboard)', 'Alta'),
    ('RF18', 'Cambiar estado de espacio', 'El administrador podrá cambiar manualmente el estado de un espacio a MANTENIMIENTO y viceversa.', 'idEspacio, espEstado', 'Media'),
    ('RF19', 'Mostrar placa en espacio ocupado', 'La vista general mostrará la placa del vehículo actualmente parqueado en espacios OCUPADOS.', 'InventarioParqueo activo (fkIdEspacio)', 'Media'),
]
for args in rfs_espacios:
    tabla_rf(*args)

h2('Módulo Inventario Parqueos')
rfs_parqueos = [
    ('RF20', 'Registrar ingreso de vehículo', 'El vigilante registrará el ingreso, seleccionando placa y espacio. El sistema registra la hora de entrada y cambia el espacio a OCUPADO.', 'fkIdVehiculo, fkIdEspacio, parHoraEntrada', 'Alta'),
    ('RF21', 'Registrar salida de vehículo', 'El vigilante registrará la salida. El sistema calcula el costo, genera el pago y libera el espacio a DISPONIBLE.', 'idParqueo, parHoraSalida, fkIdTarifa', 'Alta'),
    ('RF22', 'Calcular costo de parqueo', 'Fórmula: ceil((minutos / 60) × precioHora / 100) × 100. Visitantes usan precioHoraVisitante. Mínimo 100 COP.', 'parHoraEntrada, parHoraSalida, precioHora', 'Alta'),
    ('RF23', 'Control de espacio al ingresar/salir', 'Al ingresar: espacio → OCUPADO. Al salir: espacio → DISPONIBLE. También aplica al confirmar/cancelar reservas.', 'espEstado (ENUM)', 'Alta'),
    ('RF24', 'Ver historial de parqueos', 'Admin ve historial completo con filtros. El cliente ve solo sus propios registros.', 'Fecha, fkIdVehiculo, fkIdEspacio', 'Media'),
    ('RF25', 'Ver parqueos activos en tiempo real', 'El dashboard del vigilante muestra vehículos parqueados con espacio, hora y tiempo acumulado. AJAX cada 30 s.', 'parHoraSalida IS NULL', 'Alta'),
    ('RF26', 'Evitar doble ingreso activo', 'El sistema valida que un vehículo no pueda tener más de un registro activo simultáneo.', 'fkIdVehiculo, parHoraSalida__isnull', 'Alta'),
]
for args in rfs_parqueos:
    tabla_rf(*args)

h2('Módulo Tarifas')
rfs_tarifas = [
    ('RF27', 'Crear tarifa', 'El administrador creará tarifas con nombre, tipo de espacio, precio por hora (usuario/visitante), precio por día, precio mensual y fechas de vigencia.', 'nombre, fkIdTipoEspacio, precioHora, precioHoraVisitante, precioDia, precioMensual, fechaInicio, fechaFin', 'Alta'),
    ('RF28', 'Restricción de tarifa activa única', 'Solo una tarifa activa por tipo de espacio. Al activar una nueva, se desactiva automáticamente la anterior del mismo tipo.', 'activa (BooleanField), fkIdTipoEspacio', 'Alta'),
    ('RF29', 'Editar y desactivar tarifa', 'El administrador podrá editar o desactivar tarifas. No se eliminan tarifas con parqueos asociados.', 'idTarifa, activa', 'Media'),
    ('RF30', 'Consultar tarifa vigente', 'El sistema mostrará la tarifa activa del tipo de espacio al registrar un ingreso o calcular un costo.', 'fkIdTipoEspacio, activa=True', 'Alta'),
]
for args in rfs_tarifas:
    tabla_rf(*args)

h2('Módulo Pagos')
rfs_pagos = [
    ('RF31', 'Generar pago al registrar salida', 'Al registrar la salida se generará automáticamente un registro de pago con monto calculado, método y estado PAGADO.', 'pagMonto, pagMetodo, pagEstado, fkIdParqueo', 'Alta'),
    ('RF32', 'Aplicar cupón en pago', 'El vigilante podrá ingresar un código de cupón antes de finalizar el cobro. Si es válido, el descuento se aplica al monto total.', 'cupCodigo, pagMonto', 'Alta'),
    ('RF33', 'Ver historial de pagos', 'Admin ve todos los pagos con filtros por fecha, monto, método y estado. El cliente ve solo los suyos.', 'pagFechaPago, pagEstado', 'Media'),
    ('RF34', 'Reportes de ingresos', 'El sistema generará reportes de ingresos agrupados por período, mostrando total recaudado y número de transacciones.', 'pagFechaPago, pagMonto', 'Media'),
    ('RF35', 'Exportar reportes', 'El admin podrá descargar reportes de ingresos, ocupación, reservas y cupones en PDF (reportlab) o Excel (openpyxl).', 'Filtros de reporte seleccionados', 'Media'),
]
for args in rfs_pagos:
    tabla_rf(*args)

h2('Módulo Cupones')
rfs_cupones = [
    ('RF36', 'Crear cupón', 'El administrador creará cupones con nombre, código único alfanumérico, tipo (PORCENTAJE/VALOR_FIJO), valor, descripción, fechas de vigencia y estado activo.', 'cupNombre, cupCodigo, cupTipo, cupValor, cupFechaInicio, cupFechaFin, cupActivo', 'Alta'),
    ('RF37', 'Validar cupón', 'Al aplicar un cupón se verificará: código existente, activo=True, dentro de fechas de vigencia y con usos disponibles.', 'cupCodigo, cupActivo, cupFechaFin', 'Alta'),
    ('RF38', 'Aplicar descuento', 'PORCENTAJE: descuento = cupValor% del total. VALOR_FIJO: descuento fijo en COP. El monto final no puede ser negativo.', 'cupTipo, cupValor, pagMonto', 'Alta'),
    ('RF39', 'Registrar uso de cupón', 'Cada uso de cupón se registra en CuponAplicado con pago, cupón y monto descontado.', 'fkIdPago, fkIdCupon, montoDescontado', 'Alta'),
    ('RF40', 'Desactivar cupón vencido', 'El sistema detecta cupones con cupFechaFin pasada y los marca inactivos en el momento de la validación.', 'cupFechaFin, fecha actual', 'Media'),
    ('RF41', 'Listar cupones activos', 'El administrador consultará la lista de cupones con estado, vigencia y número de usos.', 'cupActivo, cupFechaFin', 'Media'),
]
for args in rfs_cupones:
    tabla_rf(*args)

h2('Módulo Reservas')
rfs_reservas = [
    ('RF42', 'Registrar reserva', 'El cliente o administrador podrá reservar un espacio para una fecha y hora específicas, verificando que no exista traslape con otra reserva activa.', 'resFechaReserva, resHoraInicio, resHoraFin, fkIdEspacio, fkIdVehiculo', 'Alta'),
    ('RF43', 'Consultar disponibilidad para reserva', 'El sistema mostrará espacios disponibles para una fecha y rango horario, excluyendo los que tengan reservas solapadas o estén OCUPADOS/MANTENIMIENTO.', 'resFechaReserva, resHoraInicio, resHoraFin', 'Alta'),
    ('RF44', 'Cancelar reserva', 'El cliente podrá cancelar una reserva propia siempre que no haya comenzado. Al cancelar, el espacio regresa a DISPONIBLE.', 'idReserva, resEstado', 'Alta'),
    ('RF45', 'Modificar reserva', 'El administrador o cliente podrán editar fecha u hora de una reserva, siempre que el nuevo horario esté disponible.', 'idReserva, resFechaReserva, resHoraInicio, resHoraFin', 'Media'),
    ('RF46', 'Bloquear espacio al confirmar reserva', 'Al confirmar: espacio → RESERVADO. Al cancelar: espacio → DISPONIBLE. Al convertir a parqueo: espacio → OCUPADO.', 'espEstado, resEstado', 'Alta'),
]
for args in rfs_reservas:
    tabla_rf(*args)

h2('Módulo NOVEDADES')
rfs_novedades = [
    ('RF47', 'Registrar novedad', 'El vigilante o administrador registrará una novedad asociándola a un vehículo y/o espacio, con descripción, foto (máx. 5 MB, JPG/PNG/GIF/WEBP) y estado PENDIENTE por defecto.', 'novDescripcion, novFoto, fkIdVehiculo, fkIdEspacio, fkIdReportador', 'Alta'),
    ('RF48', 'Actualizar estado de novedad', 'El administrador cambiará el estado entre PENDIENTE, EN_PROCESO y RESUELTO, agregando comentario de seguimiento.', 'novEstado, novComentario', 'Alta'),
    ('RF49', 'Asignar responsable', 'Al crear o editar una novedad se podrá asignar un responsable (usuario con rol ADMIN o VIGILANTE).', 'fkIdResponsable', 'Media'),
    ('RF50', 'Filtrar novedades', 'El listado se filtrará por estado y por texto en la descripción.', 'novEstado, novDescripcion__icontains', 'Media'),
    ('RF51', 'Notificar usuario por correo', 'Al crear o actualizar una novedad, el sistema enviará correo al dueño del vehículo afectado mediante SendGrid.', 'fkIdVehiculo.fkIdUsuario.usuCorreo', 'Alta'),
    ('RF52', 'Eliminar novedad', 'El administrador podrá eliminar una novedad del sistema (acción irreversible).', 'idNovedad', 'Baja'),
]
for args in rfs_novedades:
    tabla_rf(*args)

h2('Módulo Fidelización (Stickers)')
rfs_fidelidad = [
    ('RF53', 'Acumular sticker al completar parqueo', 'Al registrar la salida de un vehículo con propietario registrado y permanencia > 60 minutos, se asigna automáticamente un sticker al usuario.', 'parHoraEntrada, parHoraSalida, fkIdUsuario, permanencia > 60 min', 'Alta'),
    ('RF54', 'Ver stickers en perfil', 'El cliente verá en su perfil los stickers acumulados, la meta configurada y una barra de progreso. Cuando alcance la meta, aparecerá el botón de canje.', 'fkIdUsuario, metaStickers (ConfiguracionFidelidad)', 'Alta'),
    ('RF55', 'Reclamar bono por meta alcanzada', 'Al alcanzar la meta, el usuario reclama un cupón 100% generado automáticamente, con vigencia de 30 días y uso único.', 'stkFecha, metaStickers, cupValor=100, cupTipo=PORCENTAJE', 'Alta'),
    ('RF56', 'Configurar meta de stickers', 'El administrador modificará el número de stickers necesarios para el bono y los días de vigencia del cupón generado.', 'metaStickers, diasVencimientoBono (ConfiguracionFidelidad)', 'Media'),
]
for args in rfs_fidelidad:
    tabla_rf(*args)

h2('Módulo Reportes')
rfs_reportes = [
    ('RF57', 'Reporte de ingresos por período', 'El administrador generará reportes de ingresos filtrando por rango de fechas, con total recaudado, número de transacciones y promedio por operación.', 'pagFechaPago (rango), pagMonto', 'Alta'),
    ('RF58', 'Reporte de ocupación por espacio', 'El sistema generará un reporte de porcentaje de ocupación por espacio para identificar los de mayor y menor rotación.', 'fkIdEspacio, parHoraEntrada (rango)', 'Media'),
    ('RF59', 'Reporte de reservas', 'El admin verá el volumen de reservas por período con desglose por estado y tipo de espacio.', 'resFechaReserva (rango), resEstado', 'Media'),
    ('RF60', 'Reporte de cupones aplicados', 'Reporte de cupones utilizados con código, número de usos, monto total descontado y período de aplicación.', 'fkIdCupon, montoDescontado, pagFechaPago', 'Media'),
    ('RF61', 'Visualización con gráficas', 'El dashboard mostrará gráficas interactivas con Chart.js: barras de ingresos por día, líneas de ocupación y dona por tipo de espacio.', 'Datos de pagos e inventario agrupados', 'Media'),
    ('RF62', 'Exportar a PDF y Excel', 'Todos los reportes podrán descargarse en PDF (reportlab) o Excel (openpyxl) con encabezado, logo y fecha de generación.', 'Filtros de reporte seleccionados', 'Media'),
]
for args in rfs_reportes:
    tabla_rf(*args)

# ── 8.2 RNF ─────────────────────────────────────────────────────────────────
h2('Requerimientos No Funcionales')
p('Los requerimientos no funcionales establecen los atributos de calidad del sistema, siguiendo el modelo ISO/IEC 25010 (ISO, 2011).')
p()
rnfs = [
    ('RNF01', 'Rendimiento', 'El tiempo de respuesta para operaciones estándar no superará 3 segundos bajo condiciones normales.', 'Tiempo de respuesta HTTP < 3 s', 'Alta'),
    ('RNF02', 'Seguridad', 'HTTPS obligatorio, CSP mediante middleware, X-Frame-Options: DENY, X-Content-Type-Options, HSTS y protección contra fijación de sesión.', 'SecurityHeadersMiddleware, SECURE_HSTS_SECONDS', 'Alta'),
    ('RNF03', 'Usabilidad', 'Interfaz responsiva con Tailwind CSS. Tres paneles diferenciados por rol (morado/azul/verde). Validación en tiempo real en formularios.', 'Tailwind CSS CDN, 3 layouts diferenciados', 'Alta'),
    ('RNF04', 'Disponibilidad', '99% de uptime en horario operativo, garantizado por Render.com con reinicio automático ante fallos.', '99% uptime, auto-restart Render.com', 'Alta'),
    ('RNF05', 'Mantenibilidad', 'Código organizado en 11 apps Django con arquitectura MVT. Convenciones de nomenclatura consistentes.', 'Arquitectura MVT, 11 apps Django', 'Alta'),
    ('RNF06', 'Portabilidad', 'Desplegable en cualquier PaaS que soporte Python 3.12 y PostgreSQL sin depender de hardware específico.', 'requirements.txt, .env.example', 'Media'),
    ('RNF07', 'Escalabilidad', 'Paginación en vistas de más de 20 registros para evitar degradación de rendimiento.', 'Django Paginator', 'Media'),
    ('RNF08', 'Compatibilidad', 'Compatible con las dos últimas versiones de Chrome, Firefox y Microsoft Edge.', 'HTML5, CSS3, ES6+', 'Media'),
    ('RNF09', 'Internacionalización', 'Interfaz completamente en español, zona horaria America/Bogotá, valores monetarios en COP.', "LANGUAGE_CODE = 'es-co', TIME_ZONE = 'America/Bogota'", 'Alta'),
    ('RNF10', 'Trazabilidad', 'Todos los registros almacenarán marcas de tiempo de creación y actualización.', 'auto_now_add, auto_now en DateTimeField', 'Media'),
]
for code, nombre, desc, impl, prio in rnfs:
    tabla_rf(code, nombre, desc, impl, prio)

# ── 8.3 Normativos ───────────────────────────────────────────────────────────
h2('Requerimientos Normativos')
p('RNO01 — Ley 1581 de 2012 (Habeas Data): manejo de datos personales con consentimiento del titular (Congreso de la República de Colombia, 2012).')
p('RNO02 — Decreto 1377 de 2013: autorización expresa del titular antes de recopilar datos personales (Presidencia de la República de Colombia, 2013).')
p('RNO03 — Normas SENA ADSO: estructura y artefactos requeridos por el programa (SENA, 2023).')
p('RNO04 — APA 7ª edición: referencias bibliográficas y citas siguen el Manual APA 7.ª ed. (American Psychological Association, 2020).')

# ── 8.4 Reglas de Negocio ────────────────────────────────────────────────────
h2('Reglas del Negocio')
reglas = [
    ('RN01', 'Ingreso único activo por vehículo', 'Un vehículo no puede tener más de un registro con parHoraSalida IS NULL simultáneamente.'),
    ('RN02', 'Solo espacios DISPONIBLE aceptan ingresos', 'Espacios OCUPADOS, RESERVADOS o en MANTENIMIENTO no pueden recibir nuevos ingresos.'),
    ('RN03', 'Cálculo de costo minuto-base', 'costo = ceil((minutos / 60) × precioHora / 100) × 100. Visitantes pagan precioHoraVisitante. Mínimo 100 COP.'),
    ('RN04', 'Tarifa activa única por tipo de espacio', 'Al activar una nueva tarifa del mismo tipo, la anterior se desactiva automáticamente.'),
    ('RN05', 'Validación de cupón antes de aplicar', 'El cupón debe existir, estar activo y dentro de su fecha de vigencia.'),
    ('RN06', 'Pisos con espacios ocupados no se eliminan', 'No se puede eliminar un piso con espacios OCUPADOS, ni un espacio con parqueo activo.'),
    ('RN07', 'Gestión de estado al reservar/cancelar', 'Confirmar → RESERVADO. Cancelar → DISPONIBLE. Convertir a parqueo → OCUPADO.'),
    ('RN08', 'Token de recuperación de un solo uso', 'El token incorpora huella del hash de contraseña actual. Al cambiar la contraseña, el token queda invalidado.'),
    ('RN09', 'Criterio de acumulación de sticker', 'Sticker solo si: vehículo tiene propietario registrado Y permanencia > 60 minutos.'),
    ('RN10', 'Generación automática de bono al canjear', 'Al reclamar el bono se crea un cupón PORCENTAJE 100%, uso único, vigencia = diasVencimientoBono en ConfiguracionFidelidad.'),
]
for code, nombre, desc in reglas:
    pb(f'{code} — {nombre}')
    p(desc)
p()

# ═══════════════════════════════════════════════════════════════════════════
# 9. DISEÑO
# ═══════════════════════════════════════════════════════════════════════════
h1('Diseño del Sistema')

h2('Alternativas de Solución (Mockups)')
p('Durante la fase de diseño se elaboraron mockups de las interfaces principales del sistema: dashboard del administrador, vista general de pisos, formulario de ingreso de vehículos, panel del vigilante y perfil del cliente. Los mockups se encuentran en la carpeta "Mockups" de Drive del proyecto.')

h2('Diagrama de Casos de Uso')
h3('Actor: Administrador')
p('• Gestionar usuarios (crear, editar, activar/desactivar, eliminar).')
p('• Gestionar pisos, tipos de espacio y espacios.')
p('• Gestionar tarifas y cupones.')
p('• Ver y gestionar todas las reservas y novedades.')
p('• Generar y exportar reportes.')
p('• Configurar parámetros de fidelización.')

h3('Actor: Vigilante')
p('• Registrar ingreso y salida de vehículos.')
p('• Registrar vehículos visitantes.')
p('• Ver disponibilidad de espacios en tiempo real.')
p('• Registrar novedades operativas.')
p('• Aplicar cupones en el proceso de cobro.')

h3('Actor: Cliente')
p('• Registrar y gestionar sus vehículos.')
p('• Crear, consultar y cancelar reservas.')
p('• Ver su historial de parqueos y pagos.')
p('• Ver stickers acumulados y reclamar bono.')
p('• Recuperar su contraseña por correo.')

h2('Arquitectura del Software')
h3('Patrón Arquitectónico')
p('El sistema implementa el patrón MVT (Modelo-Vista-Plantilla) de Django:')
p('• Modelo: clases Python en models.py que mapean a tablas PostgreSQL mediante el ORM de Django.')
p('• Vista: clases heredadas de View que procesan la lógica de negocio y retornan respuestas HTTP.')
p('• Plantilla: archivos HTML con sintaxis Django Template Language + Tailwind CSS.')

h3('ORM de Django')
p(
    'El ORM (Object-Relational Mapping) integrado de Django gestiona toda la interacción con la base '
    'de datos sin escribir SQL directamente. Actúa como capa de abstracción que traduce automáticamente '
    'operaciones sobre objetos Python a sentencias SQL optimizadas. Ventajas en este proyecto:'
)
p('• Portabilidad: el mismo código funciona con PostgreSQL en producción sin modificaciones.')
p('• Migraciones automáticas: los cambios en modelos Python generan scripts DDL automáticamente.')
p('• Seguridad: escape automático de parámetros en todas las consultas (sin riesgo de SQL injection).')
p('• Consultas expresivas: p. ej. InventarioParqueo.objects.filter(parHoraSalida__isnull=True).')

h3('Estructura de Apps')
apps_desc = [
    ('usuarios', 'Modelo Usuario personalizado, autenticación por sesión propia, tres roles.'),
    ('vehiculos', 'Modelo Vehiculo con soporte para vehículos propios y visitantes.'),
    ('parqueadero', 'Piso, TipoEspacio, Espacio, InventarioParqueo. Panel admin y vigilante.'),
    ('tarifas', 'Modelo Tarifa con precios diferenciados por tipo de espacio y tipo de usuario.'),
    ('pagos', 'Modelo Pago vinculado a InventarioParqueo con cálculo y redondeo.'),
    ('cupones', 'Cupon y CuponAplicado. CRUD, validación y aplicación en el flujo de cobro.'),
    ('reservas', 'Reserva con gestión de disponibilidad y bloqueo/liberación de espacios.'),
    ('novedades', 'Novedad con foto, estados de progreso y notificación por correo (SendGrid).'),
    ('fidelidad', 'Sticker y ConfiguracionFidelidad. Acumulación automática al completar parqueos.'),
    ('multiparking', 'Configuración central: settings, urls, middleware de seguridad, email_utils.'),
]
for app, desc in apps_desc:
    pb(f'• {app}/: ')
    p(desc)

h2('Diagrama de Despliegue')
p('• PaaS: Render.com (Web Service + PostgreSQL add-on).')
p('• Servidor WSGI: Gunicorn.')
p('• Variables de entorno: .env (python-dotenv) con DATABASE_URL, SECRET_KEY, SENDGRID_API_KEY.')
p('• Archivos estáticos: Whitenoise.')
p('• CI/CD: cada git push a main desencadena redespliegue automático.')

h2('Modelo Entidad-Relación')
p('La base de datos está compuesta por 13 tablas normalizadas hasta tercera forma normal (3FN). Relaciones principales:')
relaciones = [
    'usuarios → vehiculos (FK: fkIdUsuario)',
    'usuarios → novedades (FK: fkIdReportador, fkIdResponsable)',
    'usuarios → fidelidad_stickers (FK: fkIdUsuario)',
    'vehiculos → inventario_parqueo (FK: fkIdVehiculo)',
    'vehiculos → reservas (FK: fkIdVehiculo)',
    'pisos → espacios (FK: fkIdPiso)',
    'tipos_espacio → espacios (FK: fkIdTipoEspacio)',
    'tipos_espacio → tarifas (FK: fkIdTipoEspacio)',
    'espacios → inventario_parqueo (FK: fkIdEspacio)',
    'espacios → reservas (FK: fkIdEspacio)',
    'inventario_parqueo → pagos (FK: fkIdParqueo)',
    'inventario_parqueo → fidelidad_stickers (FK: fkIdParqueo, OneToOne)',
    'pagos → cupones_aplicados (FK: fkIdPago)',
    'cupones → cupones_aplicados (FK: fkIdCupon)',
]
for r in relaciones:
    p(f'• {r}')

h2('Diccionario de Datos')

# Helper: tabla diccionario
def tabla_dd(titulo_tabla, nombre_tabla, campos):
    """campos: lista de (campo, tipo, longitud, nulo, pk_fk, descripcion)"""
    h3(titulo_tabla)
    tbl = doc.add_table(rows=1 + len(campos), cols=6)
    tbl.style = 'Table Grid'
    hdrs = ['Campo', 'Tipo', 'Long.', 'Nulo', 'PK/FK', 'Descripción']
    for i, h_text in enumerate(hdrs):
        cell = tbl.rows[0].cells[i]
        cell.text = h_text
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for r_idx, row_data in enumerate(campos):
        row = tbl.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            row.cells[c_idx].text = str(val)
    p(f'Nota. Tabla {nombre_tabla}. Elaboración propia (2026).')
    p()

tabla_dd('Tabla DD-01 — usuarios', 'usuarios', [
    ('id', 'BIGINT AUTO', '-', 'No', 'PK', 'Identificador único del usuario'),
    ('usuDocumento', 'VARCHAR', '15', 'No', '-', 'Número de documento (solo dígitos)'),
    ('usuNombre', 'VARCHAR', '50', 'No', '-', 'Nombre del usuario (solo letras)'),
    ('usuApellido', 'VARCHAR', '50', 'No', '-', 'Apellido del usuario (solo letras)'),
    ('usuCorreo', 'VARCHAR', '64', 'No', 'UNIQUE', 'Correo electrónico único'),
    ('usuTelefono', 'VARCHAR', '10', 'No', '-', 'Número de teléfono (10 dígitos)'),
    ('usuClaveHash', 'VARCHAR', '255', 'No', '-', 'Hash PBKDF2-SHA256 de la contraseña'),
    ('rolTipoRol', 'VARCHAR', '10', 'No', '-', 'Rol: ADMIN, VIGILANTE o CLIENTE'),
    ('usuEstado', 'BOOLEAN', '-', 'No', '-', 'True = activo, False = desactivado'),
    ('usuFechaRegistro', 'DATETIME', '-', 'No', '-', 'Fecha y hora de registro (auto)'),
])

tabla_dd('Tabla DD-02 — vehiculos', 'vehiculos', [
    ('id', 'BIGINT AUTO', '-', 'No', 'PK', 'Identificador único del vehículo'),
    ('vehPlaca', 'VARCHAR', '8', 'No', 'UNIQUE', 'Placa (alfanumérico + guión)'),
    ('vehTipo', 'VARCHAR', '5', 'No', '-', 'Tipo: CARRO, MOTO o BICI'),
    ('vehColor', 'VARCHAR', '20', 'No', '-', 'Color del vehículo'),
    ('vehMarca', 'VARCHAR', '30', 'No', '-', 'Marca del vehículo'),
    ('vehModelo', 'VARCHAR', '30', 'No', '-', 'Modelo del vehículo'),
    ('vehEstado', 'BOOLEAN', '-', 'No', '-', 'True = activo, False = inactivo'),
    ('fkIdUsuario', 'BIGINT', '-', 'Sí', 'FK→usuarios', 'Propietario; null si visitante'),
    ('nombre_contacto', 'VARCHAR', '50', 'Sí', '-', 'Nombre contacto (solo visitantes)'),
    ('telefono_contacto', 'VARCHAR', '10', 'Sí', '-', 'Teléfono contacto (solo visitantes)'),
])

tabla_dd('Tabla DD-03 — pisos', 'pisos', [
    ('id', 'BIGINT AUTO', '-', 'No', 'PK', 'Identificador del piso'),
    ('pisNombre', 'VARCHAR', '30', 'No', '-', 'Nombre del piso (ej.: Piso 1)'),
    ('pisEstado', 'BOOLEAN', '-', 'No', '-', 'True = activo, False = inactivo'),
])

tabla_dd('Tabla DD-04 — tipos_espacio', 'tipos_espacio', [
    ('id', 'BIGINT AUTO', '-', 'No', 'PK', 'Identificador del tipo'),
    ('nombre', 'VARCHAR', '20', 'No', 'UNIQUE', 'Nombre: Carro, Moto, Bicicleta'),
])

tabla_dd('Tabla DD-05 — espacios', 'espacios', [
    ('id', 'BIGINT AUTO', '-', 'No', 'PK', 'Identificador del espacio'),
    ('espNumero', 'VARCHAR', '10', 'No', '-', 'Número o código del espacio (ej.: A-01)'),
    ('fkIdPiso', 'BIGINT', '-', 'No', 'FK→pisos', 'Piso al que pertenece'),
    ('fkIdTipoEspacio', 'BIGINT', '-', 'No', 'FK→tipos_espacio', 'Tipo para aplicar tarifa'),
    ('espEstado', 'VARCHAR', '10', 'No', '-', 'DISPONIBLE / OCUPADO / RESERVADO / MANTENIMIENTO'),
])

tabla_dd('Tabla DD-06 — inventario_parqueo', 'inventario_parqueo', [
    ('id', 'BIGINT AUTO', '-', 'No', 'PK', 'Identificador del registro de parqueo'),
    ('parHoraEntrada', 'DATETIME', '-', 'No', '-', 'Fecha y hora de ingreso del vehículo'),
    ('parHoraSalida', 'DATETIME', '-', 'Sí', '-', 'Fecha y hora de salida; NULL = aún parqueado'),
    ('fkIdVehiculo', 'BIGINT', '-', 'No', 'FK→vehiculos', 'Vehículo que ocupa el espacio'),
    ('fkIdEspacio', 'BIGINT', '-', 'No', 'FK→espacios', 'Espacio asignado'),
])

tabla_dd('Tabla DD-07 — tarifas', 'tarifas', [
    ('id', 'BIGINT AUTO', '-', 'No', 'PK', 'Identificador de la tarifa'),
    ('nombre', 'VARCHAR', '50', 'No', '-', 'Nombre descriptivo de la tarifa'),
    ('fkIdTipoEspacio', 'BIGINT', '-', 'No', 'FK→tipos_espacio', 'Tipo de espacio al que aplica'),
    ('precioHora', 'DECIMAL(10,2)', '-', 'No', '-', 'Precio/hora para usuarios registrados (COP)'),
    ('precioHoraVisitante', 'DECIMAL(10,2)', '-', 'No', '-', 'Precio/hora para visitantes (COP)'),
    ('precioDia', 'DECIMAL(10,2)', '-', 'No', '-', 'Precio por día completo (COP)'),
    ('precioMensual', 'DECIMAL(10,2)', '-', 'No', '-', 'Precio mensual (COP)'),
    ('activa', 'BOOLEAN', '-', 'No', '-', 'True = vigente; solo una activa por tipo'),
    ('fechaInicio', 'DATE', '-', 'No', '-', 'Inicio de vigencia'),
    ('fechaFin', 'DATE', '-', 'Sí', '-', 'Fin de vigencia; null = indefinida'),
])

tabla_dd('Tabla DD-08 — pagos', 'pagos', [
    ('id', 'BIGINT AUTO', '-', 'No', 'PK', 'Identificador del pago'),
    ('pagFechaPago', 'DATETIME', '-', 'No', '-', 'Fecha y hora del pago'),
    ('pagMonto', 'DECIMAL(10,2)', '-', 'No', '-', 'Monto total cobrado (COP)'),
    ('pagMetodo', 'VARCHAR', '13', 'No', '-', 'EFECTIVO / TARJETA / TRANSFERENCIA'),
    ('pagEstado', 'VARCHAR', '9', 'No', '-', 'PAGADO o PENDIENTE'),
    ('fkIdParqueo', 'BIGINT', '-', 'No', 'FK→inventario_parqueo', 'Parqueo al que corresponde'),
])

tabla_dd('Tabla DD-09 — cupones', 'cupones', [
    ('id', 'BIGINT AUTO', '-', 'No', 'PK', 'Identificador del cupón'),
    ('cupNombre', 'VARCHAR', '50', 'No', '-', 'Nombre del cupón'),
    ('cupCodigo', 'VARCHAR', '20', 'No', 'UNIQUE', 'Código único de canje (mayúsculas + números)'),
    ('cupTipo', 'VARCHAR', '10', 'No', '-', 'PORCENTAJE o VALOR_FIJO'),
    ('cupValor', 'DECIMAL(10,2)', '-', 'No', '-', 'Valor del descuento'),
    ('cupDescripcion', 'TEXT', '-', 'No', '-', 'Descripción del cupón'),
    ('cupFechaInicio', 'DATE', '-', 'No', '-', 'Inicio de vigencia'),
    ('cupFechaFin', 'DATE', '-', 'No', '-', 'Fin de vigencia'),
    ('cupActivo', 'BOOLEAN', '-', 'No', '-', 'True = disponible para canje'),
])

tabla_dd('Tabla DD-10 — reservas', 'reservas', [
    ('id', 'BIGINT AUTO', '-', 'No', 'PK', 'Identificador de la reserva'),
    ('resFechaReserva', 'DATE', '-', 'No', '-', 'Fecha de la reserva'),
    ('resHoraInicio', 'TIME', '-', 'No', '-', 'Hora de inicio'),
    ('resHoraFin', 'TIME', '-', 'Sí', '-', 'Hora de fin'),
    ('resEstado', 'VARCHAR', '10', 'No', '-', 'CONFIRMADA / CANCELADA / COMPLETADA'),
    ('fkIdEspacio', 'BIGINT', '-', 'No', 'FK→espacios', 'Espacio reservado'),
    ('fkIdVehiculo', 'BIGINT', '-', 'No', 'FK→vehiculos', 'Vehículo que reserva'),
])

tabla_dd('Tabla DD-11 — novedades', 'novedades', [
    ('id', 'BIGINT AUTO', '-', 'No', 'PK', 'Identificador de la novedad'),
    ('novDescripcion', 'TEXT', '-', 'No', '-', 'Descripción detallada'),
    ('novFoto', 'VARCHAR', '100', 'Sí', '-', 'Ruta del archivo de imagen (máx. 5 MB)'),
    ('novEstado', 'VARCHAR', '10', 'No', '-', 'PENDIENTE / EN_PROCESO / RESUELTO'),
    ('novComentario', 'TEXT', '-', 'No', '-', 'Comentario de seguimiento'),
    ('fkIdVehiculo', 'BIGINT', '-', 'Sí', 'FK→vehiculos', 'Vehículo involucrado (opcional)'),
    ('fkIdEspacio', 'BIGINT', '-', 'Sí', 'FK→espacios', 'Espacio afectado (opcional)'),
    ('fkIdReportador', 'BIGINT', '-', 'Sí', 'FK→usuarios', 'Usuario que registra'),
    ('fkIdResponsable', 'BIGINT', '-', 'Sí', 'FK→usuarios', 'Usuario asignado para resolver'),
    ('novFechaCreacion', 'DATETIME', '-', 'No', '-', 'Fecha y hora de creación (auto)'),
    ('novFechaActualizacion', 'DATETIME', '-', 'No', '-', 'Última actualización (auto)'),
])

tabla_dd('Tabla DD-12 — fidelidad_stickers', 'fidelidad_stickers', [
    ('id', 'BIGINT AUTO', '-', 'No', 'PK', 'Identificador del sticker'),
    ('fkIdUsuario', 'BIGINT', '-', 'No', 'FK→usuarios', 'Usuario que acumula el sticker'),
    ('fkIdParqueo', 'BIGINT', '-', 'No', 'FK→inventario_parqueo (OneToOne)', 'Parqueo que originó el sticker'),
    ('stkFecha', 'DATETIME', '-', 'No', '-', 'Fecha y hora de asignación (auto)'),
])

h2('Políticas de Seguridad de los Datos')
p('• Contraseñas almacenadas con PBKDF2-SHA256. Nunca en texto plano.')
p('• Sesiones con cookie HttpOnly y regeneración de ID en cada inicio (cycle_key()) para prevenir fijación de sesión.')
p('• HTTPS obligatorio en producción (SECURE_SSL_REDIRECT, SECURE_HSTS_SECONDS).')
p('• Cabeceras HTTP: X-Frame-Options: DENY, X-Content-Type-Options: nosniff, X-XSS-Protection: 1; mode=block.')
p('• CSP mediante SecurityHeadersMiddleware personalizado.')
p('• Tokens de recuperación de contraseña de un solo uso (huella del hash actual).')
p('• Validación de archivos: solo imágenes (JPG/PNG/GIF/WEBP), máximo 5 MB.')
p('• Protección CSRF mediante CsrfViewMiddleware en todas las vistas POST.')

# ═══════════════════════════════════════════════════════════════════════════
# 10. CONSTRUCCIÓN
# ═══════════════════════════════════════════════════════════════════════════
h1('Construcción del Software')

h2('Base de Datos para el Software')
p(
    'En el entorno de producción (Render.com) el sistema usa PostgreSQL como motor de base de datos. '
    'Django gestiona la conexión mediante DATABASE_URL, parseada con dj-database-url. El ORM garantiza '
    'que el código Python sea idéntico entre entornos de desarrollo y producción, sin requerir cambios '
    'al cambiar de motor de base de datos. Las migraciones se generan automáticamente con '
    'python manage.py makemigrations y se aplican con python manage.py migrate.'
)

h2('Objetos de la Base de Datos (Procedimientos Almacenados, Vistas, Disparadores)')
p(
    'El proyecto define tres disparadores (triggers) a nivel de base de datos que automatizan cambios '
    'de estado críticos, independientemente del origen del cambio. Aunque el ORM de Django gestiona la '
    'mayor parte de la lógica en la capa de aplicación, los triggers garantizan consistencia incluso '
    'ante operaciones directas sobre la BD. En producción (PostgreSQL) se implementan como funciones PL/pgSQL.'
)

pb('Trigger 1 — Reserva a Parqueo:')
p('Al insertar un registro en reservas, el trigger crea automáticamente el registro de ingreso en inventario_parqueo.')
p('Disparador de prueba:')
p("INSERT INTO reservas (resHoraInicio, resHoraFin, fkIdEspacio, fkIdVehiculo)\nVALUES ('2025-09-10 10:00:00', '2025-09-10 12:00:00', 1, 1);")

pb('Trigger 2 — Bloquear espacio al ingresar:')
p('Al insertar un registro en inventario_parqueo, el trigger actualiza el estado del espacio a OCUPADO.')
p('Disparador de prueba:')
p("INSERT INTO parqueos (parHoraEntrada, fkIdVehiculo, fkIdTarifa, fkIdEspacio)\nVALUES (NOW(), 1, 1, 1);")

pb('Trigger 3 — Liberar espacio al salir:')
p('Al actualizar parHoraSalida en inventario_parqueo (valor no nulo), el trigger cambia el espacio a DISPONIBLE.')
p('Disparador de prueba:')
p("UPDATE parqueos SET parHoraSalida = NOW() WHERE idParqueo = 123;")

h2('Esquemas de Seguridad de los Datos')
p('Ver sección 9.8 Políticas de Seguridad de los Datos.')

h2('Codificación del Software')
h3('Estándar de Codificación')
p('• Nombres de campos: camelCase prefijado con abreviatura de tabla (usuNombre, vehPlaca, espEstado, parHoraEntrada).')
p('• Claves foráneas: fkId<Entidad> con db_column explícito.')
p('• db_table: declarado explícitamente en la clase Meta de cada modelo.')
p('• Vistas: CBV (class-based), patrón List/Create/Update/Delete.')
p('• Sin Django Forms: lectura directa de request.POST con validación vía re.match() y mensajes Django.')
p('• @property para campos derivados (es_visitante, usuNombreCompleto). No generan columnas en BD.')

h3('Código Fuente de Módulos — Fragmentos Representativos')
pb('Cálculo de costo de parqueo (parqueadero/vigilante_views.py):')
p("import math\nfrom decimal import Decimal\n\ndef calcular_costo(hora_entrada, hora_salida, precio_hora):\n    delta = hora_salida - hora_entrada\n    minutos = delta.total_seconds() / 60\n    costo = Decimal(str(math.ceil((minutos / 60) * float(precio_hora) / 100) * 100))\n    return max(costo, Decimal('100'))")

pb('Middleware de seguridad (multiparking/middleware.py):')
p("class SecurityHeadersMiddleware:\n    CSP = \"default-src 'self'; script-src 'self' 'unsafe-inline' https://cdn.tailwindcss.com; ...\"\n    def __call__(self, request):\n        response = self.get_response(request)\n        response['Content-Security-Policy'] = self.CSP\n        response['X-XSS-Protection'] = '1; mode=block'\n        return response")

h2('Servicios Web')
p('El framework Django REST Framework (DRF) está instalado en el proyecto. Los endpoints se organizan bajo /api/<app>/. En la versión actual se exponen endpoints de consulta para espacios, reservas y parqueos activos, consumidos principalmente por el dashboard AJAX del vigilante.')

h2('Control de Versiones')
p('El proyecto utiliza Git como sistema de control de versiones, con repositorio en GitHub (rama main). El flujo de trabajo consiste en commits frecuentes con mensajes descriptivos. Los archivos sensibles (.env, __pycache__) están excluidos mediante .gitignore. Cada push a main desencadena redespliegue automático en Render.com.')

# ═══════════════════════════════════════════════════════════════════════════
# 11. PRUEBAS
# ═══════════════════════════════════════════════════════════════════════════
h1('Pruebas del Software')

h2('Planeación y Diseño de Pruebas a Nivel Unitario')
p('Las pruebas unitarias verifican el correcto funcionamiento de cada módulo de forma independiente:')

tabla_generica(
    ['ID', 'Módulo', 'Caso de Prueba', 'Resultado Esperado', 'Estado'],
    [
        ['PU01', 'Usuarios', 'Registro con datos válidos', 'Usuario creado, contraseña hasheada', 'PASÓ'],
        ['PU02', 'Usuarios', 'Login con credenciales correctas', 'Sesión creada con rol CLIENTE', 'PASÓ'],
        ['PU03', 'Usuarios', 'Login con contraseña incorrecta', 'Error "Credenciales incorrectas"', 'PASÓ'],
        ['PU04', 'Usuarios', 'Recuperación de contraseña', 'Correo enviado, token generado', 'PASÓ'],
        ['PU05', 'Usuarios', 'Token usado segunda vez', 'Error "enlace ya utilizado"', 'PASÓ'],
        ['PU06', 'Vehículos', 'Registro vehículo propio', 'Vehículo asociado al usuario', 'PASÓ'],
        ['PU07', 'Vehículos', 'Registro vehículo visitante', 'Vehículo sin fkIdUsuario', 'PASÓ'],
        ['PU08', 'Espacios', 'Crear espacio en piso activo', 'Espacio en estado DISPONIBLE', 'PASÓ'],
        ['PU09', 'Parqueos', 'Ingreso en espacio DISPONIBLE', 'Estado cambia a OCUPADO', 'PASÓ'],
        ['PU10', 'Parqueos', 'Ingreso con vehículo ya activo', 'Error de doble ingreso', 'PASÓ'],
        ['PU11', 'Parqueos', 'Calcular costo 90 min a $5000/h', '$8000 (ceil a 100)', 'PASÓ'],
        ['PU12', 'Tarifas', 'Activar nueva tarifa del mismo tipo', 'Solo una activa por tipo', 'PASÓ'],
        ['PU13', 'Cupones', 'Cupón PORCENTAJE 20% en $10000', 'Descuento $2000, total $8000', 'PASÓ'],
        ['PU14', 'Cupones', 'Aplicar cupón vencido', 'Error "cupón vencido"', 'PASÓ'],
        ['PU15', 'Reservas', 'Crear reserva en espacio disponible', 'Espacio → RESERVADO', 'PASÓ'],
        ['PU16', 'Reservas', 'Cancelar reserva', 'Espacio → DISPONIBLE', 'PASÓ'],
        ['PU17', 'Novedades', 'Registrar novedad con foto PNG', 'Novedad guardada, foto almacenada', 'PASÓ'],
        ['PU18', 'Novedades', 'Subir archivo .exe como foto', 'Error "Solo se permiten imágenes"', 'PASÓ'],
        ['PU19', 'Fidelidad', 'Parqueo >60 min con usuario registrado', 'Sticker asignado al usuario', 'PASÓ'],
        ['PU20', 'Fidelidad', 'Parqueo <60 min', 'Sin sticker asignado', 'PASÓ'],
    ],
    titulo='Tabla PT-01\nPruebas Unitarias por Módulo'
)
p('Nota. Pruebas ejecutadas manualmente en entorno de desarrollo (2026).')

h2('Pruebas a Nivel Sistema')
tabla_generica(
    ['ID', 'Flujo', 'Descripción', 'Estado'],
    [
        ['PS01', 'Ingreso y salida completo', 'Admin crea espacio → Vigilante registra ingreso → OCUPADO → Salida → Pago generado → DISPONIBLE', 'PASÓ'],
        ['PS02', 'Reserva por cliente', 'Cliente crea reserva → RESERVADO → Vigilante convierte a parqueo → OCUPADO → Salida → Pago', 'PASÓ'],
        ['PS03', 'Cupón en cobro', 'Vigilante registra salida → Ingresa cupón → Descuento aplicado → CuponAplicado registrado', 'PASÓ'],
        ['PS04', 'Novedad con notificación', 'Vigilante crea novedad con foto → Correo enviado al dueño → Admin actualiza estado → Nuevo correo', 'PASÓ'],
        ['PS05', 'Fidelización completa', 'Parqueo >1h → Sticker asignado → Cliente ve progreso → Alcanza meta → Reclama → Cupón 100% generado', 'PASÓ'],
        ['PS06', 'Recuperación de contraseña', 'Solicita correo → Recibe enlace → Nueva contraseña → Token invalidado → Reintento falla', 'PASÓ'],
        ['PS07', 'Acceso sin autenticación', 'Usuario no autenticado accede a /admin-panel/ → Redirige a /login/ (302)', 'PASÓ'],
        ['PS08', 'Acceso con rol incorrecto', 'CLIENTE accede a /admin-panel/ → Redirigido a /dashboard/ (acceso denegado)', 'PASÓ'],
    ],
    titulo='Tabla PT-02\nPruebas de Sistema — Flujos End-to-End'
)
p('Nota. Elaboración propia (2026).')

h2('Manejo de Alertas / Excepciones')
p('• Errores de validación: mensajes de error mediante el framework messages de Django con estilos Tailwind.')
p('• Errores 404: página personalizada sin exponer información del stack.')
p('• Errores de base de datos: capturados con try/except con mensajes amigables.')
p('• Errores de correo: capturados sin interrumpir el flujo principal; error registrado en log.')
p('• Archivos inválidos: tipo MIME y tamaño verificados antes de guardar.')

# ═══════════════════════════════════════════════════════════════════════════
# 12. DESPLIEGUE
# ═══════════════════════════════════════════════════════════════════════════
h1('Plan de Despliegue')

h2('Plataforma de Despliegue')
p('El sistema MultiParking se despliega en Render.com (Web Service + PostgreSQL add-on). El proceso CI/CD es automático: cada push a main desencadena un redespliegue.')

h2('Configuración de Variables de Entorno')
p('SECRET_KEY=<clave secreta Django 50+ caracteres>')
p('DEBUG=False')
p('DATABASE_URL=<URL de conexión PostgreSQL>')
p('SENDGRID_API_KEY=<clave API SendGrid>')
p('DEFAULT_FROM_EMAIL=no-reply@multiparking.com')
p('ALLOWED_HOSTS=multiparking.onrender.com')

h2('Proceso de Despliegue')
p('1. git push origin main → Render detecta el cambio e inicia el build.')
p('2. pip install -r requirements.txt → instalación de dependencias.')
p('3. python manage.py collectstatic --noinput → archivos estáticos con Whitenoise.')
p('4. python manage.py migrate → migraciones pendientes.')
p('5. gunicorn multiparking.wsgi:application --bind 0.0.0.0:$PORT → servidor.')

h2('Manual del Usuario')
h3('Rol Administrador (/admin-panel/)')
p('• Gestión de usuarios: crear, editar, activar/desactivar.')
p('• Gestión de espacios: configurar pisos, tipos y espacios.')
p('• Gestión de tarifas: definir precios por tipo de espacio.')
p('• Gestión de cupones: crear y administrar descuentos.')
p('• Reportes: visualizar estadísticas y exportar a PDF/Excel.')
p('• Novedades: gestionar incidentes operativos.')

h3('Rol Vigilante (/vigilante/dashboard/)')
p('• Registrar ingreso: seleccionar vehículo y espacio. Sistema registra hora y marca OCUPADO.')
p('• Registrar salida: seleccionar registro activo, aplicar cupón opcional, confirmar monto y método de pago.')
p('• Vista de pisos: visualización en tiempo real con actualización cada 30 segundos.')
p('• Registrar novedad: documentar incidentes con descripción, foto, espacio y vehículo.')

h3('Rol Cliente (/dashboard/)')
p('• Mis vehículos: registrar y gestionar vehículos propios.')
p('• Mis reservas: crear, consultar y cancelar reservas anticipadas.')
p('• Mi historial: consultar parqueos y pagos anteriores.')
p('• Mi perfil: ver stickers acumulados, barra de progreso y canjear bono al alcanzar la meta.')

h2('Copias de Seguridad de Datos y Respaldos')
p('Render.com provee snapshots automáticos diarios de la base de datos PostgreSQL. Se recomienda programar dumps adicionales con:')
p('pg_dump -U <usuario> <base_de_datos> > backup_$(date +%Y%m%d).sql')

# ═══════════════════════════════════════════════════════════════════════════
# 13. CONCLUSIONES
# ═══════════════════════════════════════════════════════════════════════════
h1('Conclusiones')
p(
    'El desarrollo del sistema de información web MultiParking permitió automatizar de forma integral '
    'los procesos operativos de un parqueadero privado, sustituyendo métodos manuales por una plataforma '
    'digital escalable, segura y de fácil uso. Se validaron los conceptos fundamentales de análisis y '
    'desarrollo de software: levantamiento de requerimientos, diseño de base de datos normalizada, '
    'implementación con patrón MVT, pruebas funcionales y despliegue en la nube.'
)
p(
    'La elección de Django 6.0 demostró ser acertada: su sistema de migraciones facilitó la evolución '
    'incremental del modelo de datos, su ORM redujo la complejidad del acceso a PostgreSQL y su arquitectura '
    'de apps permitió organizar el código de forma modular. La autenticación propia sin django.contrib.auth.User '
    'fortaleció la comprensión de los mecanismos de seguridad en sesiones web.'
)
p(
    'El módulo de fidelización y el sistema de novedades con notificación por correo aportaron valor '
    'diferencial a la solución. Como trabajo futuro se identifican: integración con hardware automatizado, '
    'aplicación móvil nativa e implementación de pagos digitales en línea (PayU, Nequi).'
)

# ═══════════════════════════════════════════════════════════════════════════
# 14. BIBLIOGRAFÍA
# ═══════════════════════════════════════════════════════════════════════════
h1('Bibliografía y Ciberografía')
referencias = [
    'American Psychological Association. (2020). Publication manual of the American Psychological Association (7th ed.). https://doi.org/10.1037/0000165-000',
    'Congreso de la República de Colombia. (2012). Ley 1581 de 2012: por la cual se dictan disposiciones generales para la protección de datos personales. Diario Oficial, 48587.',
    'Django Software Foundation. (2024). Django documentation (Version 6.0). https://docs.djangoproject.com/en/6.0/',
    'García, A., & Martínez, L. (2022). Implementación de sistemas de información para la gestión de parqueaderos en ciudades intermedias de Colombia. Revista Ingeniería e Innovación, 10(2), 45–58.',
    'Garavito Albao, G., Macías Lamprea, R. A., & Pinzón Osorio, N. (2021). Sistema de gestión para ópticas SIGEOP. TIA — Tecnología, Investigación y Academia, 9(2).',
    'ISO/IEC. (2011). ISO/IEC 25010:2011 — Systems and software engineering: Systems and software quality requirements and evaluation (SQuaRE). International Organization for Standardization.',
    'MinTIC. (2023). Política de transformación digital e inteligencia artificial de Colombia. https://www.mintic.gov.co',
    'PostgreSQL Global Development Group. (2024). PostgreSQL 16 documentation. https://www.postgresql.org/docs/16/',
    'Presidencia de la República de Colombia. (2013). Decreto 1377 de 2013: por el cual se reglamenta parcialmente la Ley 1581 de 2012. Diario Oficial, 48834.',
    'Render. (2024). Render documentation. https://render.com/docs',
    'SendGrid. (2024). SendGrid email API documentation. https://docs.sendgrid.com/',
    'Tailwind CSS. (2024). Tailwind CSS documentation (v3). https://tailwindcss.com/docs',
]
for ref in referencias:
    p(ref)

# ═══════════════════════════════════════════════════════════════════════════
# 15. GLOSARIO
# ═══════════════════════════════════════════════════════════════════════════
h1('Glosario')
glosario = [
    ('CSP (Content Security Policy)', 'Política de seguridad HTTP que controla qué recursos puede cargar el navegador, reduciendo el riesgo de ataques XSS.'),
    ('Cupón', 'Código alfanumérico que otorga un descuento (porcentaje o valor fijo) al aplicarse en el cobro del parqueo.'),
    ('Django', 'Framework web de alto nivel basado en Python que sigue el patrón arquitectónico MVT (Modelo-Vista-Plantilla).'),
    ('Espacio', 'Unidad física del parqueadero, identificada por número y ubicada en un piso. Estados: DISPONIBLE, OCUPADO, RESERVADO, MANTENIMIENTO.'),
    ('Fidelización', 'Sistema de stickers que recompensa a usuarios frecuentes con cupones de descuento automáticos al alcanzar una meta.'),
    ('Hash PBKDF2', 'Algoritmo de derivación de claves que convierte la contraseña en representación irreversible, protegiendo credenciales ante filtraciones.'),
    ('Inventario de Parqueo', 'Registro de cada entrada y salida de vehículo, incluyendo horas, espacio asignado y referencia al pago generado.'),
    ('ORM', 'Object-Relational Mapping. Capa de abstracción que permite interactuar con la base de datos usando objetos Python en lugar de SQL directo.'),
    ('PostgreSQL', 'Sistema de gestión de bases de datos relacional de código abierto utilizado en producción para almacenar los datos de MultiParking.'),
    ('Novedad', 'Registro de un incidente, daño u observación operativa con seguimiento de estado y notificación al usuario afectado.'),
    ('Parqueadero', 'Establecimiento destinado al estacionamiento de vehículos, compuesto por pisos y espacios administrados por el sistema.'),
    ('Render.com', 'Plataforma cloud de despliegue continuo (PaaS) utilizada para alojar la aplicación web y la base de datos PostgreSQL.'),
    ('Reserva', 'Apartado anticipado de un espacio para una fecha y horario específicos, realizado por un cliente registrado.'),
    ('SendGrid', 'Servicio de correo electrónico transaccional para notificaciones de recuperación de contraseña y alertas de novedades.'),
    ('Sticker', 'Unidad de fidelización acumulable otorgada al usuario cada vez que completa un parqueo de más de 60 minutos.'),
    ('Tailwind CSS', 'Framework CSS de clases utilitarias para construir interfaces responsivas directamente en el HTML.'),
    ('Tarifa', 'Precio por tipo de espacio y tipo de usuario (registrado/visitante), expresado en COP por hora, día o mes.'),
    ('Vigilante', 'Usuario con rol VIGILANTE que opera el control de ingreso y salida de vehículos, registra novedades y aplica cupones.'),
]
for term, defi in glosario:
    p()
    pb(term)
    p(defi)

# ═══════════════════════════════════════════════════════════════════════════
# GUARDAR
# ═══════════════════════════════════════════════════════════════════════════
out_path = 'Documentacion_Multiparking_COMPLETA.docx'
doc.save(out_path)
print(f'Guardado: {out_path}')
