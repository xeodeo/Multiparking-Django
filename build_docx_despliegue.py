#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Convierte manual/despliegue/manual-despliegue.md a manual/1787097177998.docx
Reutiliza el mismo enfoque de build_docx_manuales.py (Markdown -> docx a mano).
"""
import os
import re
import sys

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

sys.stdout.reconfigure(encoding='utf-8')

FONT = 'Calibri'
COLOR_H1 = '1F4E79'
COLOR_H2 = '2E6DB4'
COLOR_H3 = '3B7DBF'
COLOR_BODY = '1A1A1A'

INLINE_RE = re.compile(r'(\*\*.+?\*\*|\*.+?\*|`.+?`|\[.+?\]\(.+?\))')


def add_inline_runs(paragraph, text, base_size=10.5, base_color=COLOR_BODY):
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            _run(paragraph, text[pos:m.start()], base_size, base_color)
        token = m.group(0)
        if token.startswith('**'):
            _run(paragraph, token[2:-2], base_size, base_color, bold=True)
        elif token.startswith('`'):
            _run(paragraph, token[1:-1], base_size, base_color, mono=True)
        elif token.startswith('['):
            mm = re.match(r'\[(.+?)\]\((.+?)\)', token)
            _run(paragraph, mm.group(1), base_size, '2E6DB4', underline=True)
        elif token.startswith('*'):
            _run(paragraph, token[1:-1], base_size, base_color, italic=True)
        pos = m.end()
    if pos < len(text):
        _run(paragraph, text[pos:], base_size, base_color)


def _run(paragraph, text, size, color, bold=False, italic=False, underline=False, mono=False):
    if not text:
        return
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Consolas' if mono else FONT
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level):
    sizes = {1: 22, 2: 15, 3: 12.5, 4: 11}
    colors = {1: COLOR_H1, 2: COLOR_H2, 3: COLOR_H3, 4: COLOR_H3}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 14)
    p.paragraph_format.space_after = Pt(10 if level == 1 else 6)
    if level == 1:
        p.paragraph_format.page_break_before = False
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(sizes.get(level, 11))
    run.font.name = FONT
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(colors.get(level, COLOR_BODY))
    return p


def add_paragraph_text(doc, text, size=10.5, center=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline_runs(p, text, base_size=size)
    return p


def add_list_item(doc, text, ordered, indent=0):
    style = 'List Number' if ordered else 'List Bullet'
    try:
        p = doc.add_paragraph(style=style)
    except KeyError:
        p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25 + indent * 0.25)
    p.paragraph_format.space_after = Pt(3)
    add_inline_runs(p, text)
    return p


def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_image(doc, base_dir, rel_path, caption=None):
    full_path = os.path.join(base_dir, rel_path)
    if not os.path.exists(full_path):
        add_paragraph_text(doc, f'[Imagen no encontrada: {rel_path}]')
        return
    with Image.open(full_path) as im:
        w, h = im.size
    max_width_in = 5.8
    width_in = min(max_width_in, w / 150)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(full_path, width=Inches(width_in))
    p.paragraph_format.space_after = Pt(4)
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline_runs(cap, caption, base_size=9, base_color='555555')
        cap.paragraph_format.space_after = Pt(10)


def add_table(doc, rows):
    n_cols = len(rows[0])
    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.style = 'Light Grid Accent 1'
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx >= n_cols:
                continue
            cell = tbl.rows[r_idx].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            add_inline_runs(p, cell_text.strip(), base_size=9.5)
            for run in p.runs:
                if r_idx == 0:
                    run.font.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def parse_table_row(line):
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    return [c for c in line.split('|')]


def is_table_separator(line):
    return bool(re.match(r'^\s*\|?[\s:|-]+\|?\s*$', line)) and '-' in line


def convert(md_path, docx_path):
    base_dir = os.path.dirname(md_path)
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.read().split('\n')

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for m in ('left_margin', 'right_margin', 'top_margin', 'bottom_margin'):
        setattr(section, m, Inches(0.9))

    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(10.5)

    i = 0
    n = len(lines)
    cover_mode = True  # centrar párrafos de portada hasta el primer '---'

    while i < n:
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == '---':
            if cover_mode:
                cover_mode = False
                i += 1
                continue
            add_hr(doc)
            i += 1
            continue

        m = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if m:
            level = min(len(m.group(1)), 4)
            text = re.sub(r'\s*\{#.*?\}\s*$', '', m.group(2)).strip()
            add_heading(doc, text, level)
            i += 1
            continue

        m = re.match(r'^!\[(.*?)\]\((.*?)\)$', stripped)
        if m:
            alt, path = m.group(1), m.group(2)
            caption = None
            if i + 1 < n and lines[i + 1].strip().startswith('*') and lines[i + 1].strip().endswith('*'):
                caption = lines[i + 1].strip()[1:-1]
                i += 1
            add_image(doc, base_dir, path, caption or alt)
            i += 1
            continue

        if stripped.startswith('|'):
            table_lines = []
            while i < n and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            rows = [parse_table_row(l) for l in table_lines if not is_table_separator(l)]
            if rows:
                add_table(doc, rows)
            continue

        if stripped.startswith('>'):
            add_paragraph_text(doc, stripped.lstrip('> ').strip())
            i += 1
            continue

        m = re.match(r'^(\s*)[-*]\s+(.*)$', line)
        if m:
            indent = len(m.group(1)) // 2
            add_list_item(doc, m.group(2), ordered=False, indent=indent)
            i += 1
            continue

        m = re.match(r'^(\s*)\d+\.\s+(.*)$', line)
        if m:
            indent = len(m.group(1)) // 2
            add_list_item(doc, m.group(2), ordered=True, indent=indent)
            i += 1
            continue

        buf = [stripped]
        i += 1
        while i < n and lines[i].strip() and not re.match(r'^(#{1,6}\s|!\[|\||>|---\s*$|\s*[-*]\s|\s*\d+\.\s)', lines[i]):
            buf.append(lines[i].strip())
            i += 1
        add_paragraph_text(doc, ' '.join(buf), center=cover_mode)

    doc.save(docx_path)
    print(f'Guardado: {docx_path}')


if __name__ == '__main__':
    root = os.path.dirname(os.path.abspath(__file__))
    convert(
        os.path.join(root, 'manual/despliegue/manual-despliegue.md'),
        os.path.join(root, 'manual/1787097177998.docx'),
    )
