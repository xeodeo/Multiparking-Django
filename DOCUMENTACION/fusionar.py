"""
Fusiona los documentos de secciones en MultiParking_Documentacion_FUSIONADO.docx
  1. MultiParking_Documentacion.docx  (portada, resumen, secciones 1-4)
  2. Seccion_05_Requerimientos.docx
  3. Seccion_06_Analisis_Especificacion.docx
"""
import os
import shutil
from docxcompose.composer import Composer
from docx import Document

BASE = os.path.dirname(os.path.abspath(__file__))
SECCIONES = os.path.join(BASE, "Secciones")

DOC_BASE = os.path.join(BASE, "MultiParking_Documentacion.docx")
DOC_05   = os.path.join(SECCIONES, "Seccion_05_Requerimientos", "Seccion_05_Requerimientos.docx")
DOC_06   = os.path.join(SECCIONES, "Seccion_06_Analisis_Especificacion", "Seccion_06_Analisis_Especificacion.docx")
OUT      = os.path.join(BASE, "MultiParking_Documentacion_FUSIONADO.docx")
BAK      = OUT + ".bak"

for f in [DOC_BASE, DOC_05, DOC_06]:
    if not os.path.exists(f):
        print(f"ERROR: no existe {f}")
        raise SystemExit(1)

# Backup
shutil.copy2(OUT, BAK)
print(f"Backup -> {BAK}")

# Componer
master = Document(DOC_BASE)
composer = Composer(master)

print(f"Agregando sección 05: {DOC_05}")
composer.append(Document(DOC_05))

print(f"Agregando sección 06: {DOC_06}")
composer.append(Document(DOC_06))

composer.save(OUT)
size_kb = os.path.getsize(OUT) / 1024
print(f"\nOK -> {OUT}  ({size_kb:.1f} KB)")
